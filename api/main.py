from __future__ import annotations

import json
import os
import hashlib
import io
import secrets
import shutil
import sqlite3
import unicodedata
import uuid
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import FastAPI, Header, HTTPException, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


BASE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BASE_DIR.parent
LEGACY_DATA_DIR = BASE_DIR / "data"
DATA_DIR = Path(os.getenv("ATLASFLOW_DATA_DIR", PROJECT_ROOT / "data"))
DB_PATH = Path(os.getenv("ATLASFLOW_DB", DATA_DIR / "atlasflow.sqlite3"))
REQUIRE_PASSWORDS = os.getenv("ATLASFLOW_REQUIRE_PASSWORDS", "").strip().lower() in {"1", "true", "yes", "sim"}

STATUSES = {"entrada", "analisar", "criar", "revisar", "devolver", "concluido"}
PRIORITIES = {"normal", "urgente"}
SESSIONS: dict[str, str] = {}


class UserCreate(BaseModel):
    name: str = Field(..., min_length=1)
    email: str = ""
    department: str = ""
    role: str = "Operador"
    password: str = ""


class User(BaseModel):
    id: str
    name: str
    email: str = ""
    department: str = ""
    role: str = "Operador"
    password_enabled: bool = False
    created_at: str


class UserPasswordUpdate(BaseModel):
    admin_user_id: str
    password: str = ""


class AuthLogin(BaseModel):
    user_id: str
    password: str


class AuthSession(BaseModel):
    user_id: str
    name: str
    department: str = ""
    role: str = ""
    token: str


class ImportPayload(BaseModel):
    users: list[dict[str, Any]] = Field(default_factory=list)
    processes: list[dict[str, Any]] = Field(default_factory=list)
    currentUserId: str = ""
    current_user_id: str = ""


class ImportResult(BaseModel):
    users: int
    processes: int
    history: int
    current_user_id: str = ""


class DocumentExport(BaseModel):
    title: str = "Documento AtlasFlow"
    text: str = Field(..., min_length=1)
    filename: str = "documento"


class ProcessBase(BaseModel):
    number: str = ""
    year: str = ""
    subject: str = Field(..., min_length=1)
    secretary: str = ""
    owner: str = ""
    priority: str = "normal"
    arrival_date: str = ""
    from_sector: str = ""
    purpose: str = ""
    deadline: str = ""
    status: str = "entrada"
    exit_date: str = ""
    to_sector: str = ""
    exit_purpose: str = ""
    docs: list[str] = Field(default_factory=list)
    notes: str = ""


class ProcessCreate(ProcessBase):
    pass


class ProcessUpdate(ProcessBase):
    pass


class Process(ProcessBase):
    id: str
    user_id: str
    created_at: str
    updated_at: str


class MovementCreate(BaseModel):
    date: str = ""
    status: str = "entrada"
    to: str = ""
    purpose: str = ""
    notes: str = ""


class HistoryItem(MovementCreate):
    id: str
    process_id: str
    user_id: str
    action: str
    created_at: str


class ProcessWithHistory(Process):
    history: list[HistoryItem] = Field(default_factory=list)


app = FastAPI(
    title="AtlasFlow API",
    description="API para gestao de usuarios, processos administrativos e historico operacional.",
    version="0.2.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@contextmanager
def db() -> Any:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def now() -> str:
    return datetime.utcnow().isoformat(timespec="seconds") + "Z"


def new_id() -> str:
    return str(uuid.uuid4())


def password_hash(password: str, salt: str) -> str:
    return hashlib.sha256(f"{salt}:{password}".encode("utf-8")).hexdigest()


def build_password_fields(password: str) -> tuple[str, str]:
    if not password:
        return "", ""
    salt = secrets.token_hex(16)
    return salt, password_hash(password, salt)


def normalize_role(role: str) -> str:
    normalized = unicodedata.normalize("NFD", role or "")
    normalized = "".join(char for char in normalized if unicodedata.category(char) != "Mn")
    return normalized.strip().lower()


def is_admin_row(user: sqlite3.Row) -> bool:
    role = normalize_role(user["role"])
    return user["id"] == "user-compras" or "administrador" in role or role == "admin"


def is_read_only_row(user: sqlite3.Row) -> bool:
    role = normalize_role(user["role"])
    return any(marker in role for marker in {"consulta", "visitante", "leitura", "somente leitura"})


def require_admin_user(conn: sqlite3.Connection, user_id: str) -> sqlite3.Row:
    admin = ensure_user(conn, user_id)
    if not is_admin_row(admin):
        raise HTTPException(status_code=403, detail="Apenas administradores podem executar esta acao")
    return admin


def bearer_token(authorization: str) -> str:
    prefix = "Bearer "
    if not authorization.startswith(prefix):
        raise HTTPException(status_code=401, detail="Sessao obrigatoria")
    return authorization[len(prefix):].strip()


def require_session_user(conn: sqlite3.Connection, authorization: str) -> sqlite3.Row:
    token = bearer_token(authorization)
    user_id = SESSIONS.get(token)
    if not user_id:
        raise HTTPException(status_code=401, detail="Sessao invalida ou expirada")
    return ensure_user(conn, user_id)


def require_user_access(conn: sqlite3.Connection, user_id: str, authorization: str) -> sqlite3.Row:
    session_user = require_session_user(conn, authorization)
    if session_user["id"] != user_id and not is_admin_row(session_user):
        raise HTTPException(status_code=403, detail="Usuario sem acesso a esta mesa")
    return session_user


def require_admin_access(conn: sqlite3.Connection, admin_user_id: str, authorization: str) -> sqlite3.Row:
    session_user = require_session_user(conn, authorization)
    if session_user["id"] != admin_user_id:
        raise HTTPException(status_code=403, detail="Sessao nao corresponde ao administrador informado")
    if not is_admin_row(session_user):
        raise HTTPException(status_code=403, detail="Apenas administradores podem executar esta acao")
    return session_user


def require_process_access(conn: sqlite3.Connection, process: sqlite3.Row, authorization: str) -> sqlite3.Row:
    return require_user_access(conn, process["user_id"], authorization)


def require_process_write(conn: sqlite3.Connection, process: sqlite3.Row, authorization: str) -> sqlite3.Row:
    session_user = require_process_access(conn, process, authorization)
    if is_read_only_row(session_user):
        raise HTTPException(status_code=403, detail="Perfil permite apenas consulta")
    return session_user


def ensure_column(conn: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    columns = {row["name"] for row in conn.execute(f"pragma table_info({table})").fetchall()}
    if column not in columns:
        conn.execute(f"alter table {table} add column {column} {definition}")


def validate_status(value: str) -> None:
    if value not in STATUSES:
        raise HTTPException(status_code=422, detail=f"Status invalido: {value}")


def validate_priority(value: str) -> None:
    if value not in PRIORITIES:
        raise HTTPException(status_code=422, detail=f"Prioridade invalida: {value}")


def row_to_user(row: sqlite3.Row) -> dict[str, Any]:
    item = dict(row)
    item["password_enabled"] = bool(item.get("password_hash"))
    item.pop("password_hash", None)
    item.pop("password_salt", None)
    return item


def row_to_process(row: sqlite3.Row) -> dict[str, Any]:
    item = dict(row)
    item["docs"] = json.loads(item.get("docs") or "[]")
    return item


def row_to_history(row: sqlite3.Row) -> dict[str, Any]:
    return dict(row)


def data_get(data: dict[str, Any], *names: str, default: Any = "") -> Any:
    for name in names:
        if name in data and data[name] is not None:
            return data[name]
    return default


def as_docs(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value]
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            if isinstance(parsed, list):
                return [str(item) for item in parsed]
        except json.JSONDecodeError:
            return [item.strip() for item in value.split(",") if item.strip()]
    return []


def ensure_user(conn: sqlite3.Connection, user_id: str) -> sqlite3.Row:
    row = conn.execute("select * from users where id = ?", (user_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Usuario nao encontrado")
    return row


def ensure_process(conn: sqlite3.Connection, process_id: str) -> sqlite3.Row:
    row = conn.execute("select * from processes where id = ?", (process_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Processo nao encontrado")
    return row


def insert_history(
    conn: sqlite3.Connection,
    *,
    process_id: str,
    user_id: str,
    date: str,
    action: str,
    status_value: str,
    to: str,
    purpose: str,
    notes: str,
) -> None:
    conn.execute(
        """
        insert into history (id, process_id, user_id, date, action, status, to_sector, purpose, notes, created_at)
        values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (new_id(), process_id, user_id, date or now()[:10], action, status_value, to, purpose, notes, now()),
    )


def upsert_import_user(conn: sqlite3.Connection, item: dict[str, Any]) -> str:
    user_id = str(data_get(item, "id", default=new_id()))
    name = str(data_get(item, "name", default="Usuario"))
    email = str(data_get(item, "email", default=""))
    department = str(data_get(item, "department", default=""))
    role = str(data_get(item, "role", default="Operador"))
    created_at = str(data_get(item, "createdAt", "created_at", default=now()))
    password_hash_value = str(data_get(item, "password_hash", default=""))
    password_salt_value = str(data_get(item, "password_salt", default=""))

    conn.execute(
        """
        insert into users (id, name, email, department, role, password_hash, password_salt, created_at)
        values (?, ?, ?, ?, ?, ?, ?, ?)
        on conflict(id) do update set
            name = excluded.name,
            email = excluded.email,
            department = excluded.department,
            role = excluded.role
        """,
        (user_id, name, email, department, role, password_hash_value, password_salt_value, created_at),
    )
    return user_id


def upsert_import_process(conn: sqlite3.Connection, item: dict[str, Any], fallback_user_id: str) -> tuple[str, int]:
    process_id = str(data_get(item, "id", default=new_id()))
    user_id = str(data_get(item, "userId", "user_id", default=fallback_user_id))
    if not conn.execute("select 1 from users where id = ?", (user_id,)).fetchone():
        upsert_import_user(conn, {"id": user_id, "name": "Usuario importado"})

    status_value = str(data_get(item, "status", default="entrada"))
    priority_value = str(data_get(item, "priority", default="normal"))
    if status_value not in STATUSES:
        status_value = "entrada"
    if priority_value not in PRIORITIES:
        priority_value = "normal"

    created_at = str(data_get(item, "createdAt", "created_at", default=now()))
    updated_at = str(data_get(item, "updatedAt", "updated_at", default=created_at))
    docs_json = json.dumps(as_docs(data_get(item, "docs", default=[])), ensure_ascii=False)

    values = (
        process_id,
        user_id,
        str(data_get(item, "number", default="")),
        str(data_get(item, "year", default="")),
        str(data_get(item, "subject", default="Sem objeto")),
        str(data_get(item, "secretary", default="")),
        str(data_get(item, "owner", default="")),
        priority_value,
        str(data_get(item, "arrivalDate", "arrival_date", default="")),
        str(data_get(item, "fromSector", "from_sector", default="")),
        str(data_get(item, "purpose", default="")),
        str(data_get(item, "deadline", default="")),
        status_value,
        str(data_get(item, "exitDate", "exit_date", default="")),
        str(data_get(item, "toSector", "to_sector", default="")),
        str(data_get(item, "exitPurpose", "exit_purpose", default="")),
        docs_json,
        str(data_get(item, "notes", default="")),
        created_at,
        updated_at,
    )
    conn.execute(
        """
        insert into processes (
            id, user_id, number, year, subject, secretary, owner, priority, arrival_date,
            from_sector, purpose, deadline, status, exit_date, to_sector, exit_purpose,
            docs, notes, created_at, updated_at
        )
        values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        on conflict(id) do update set
            user_id = excluded.user_id,
            number = excluded.number,
            year = excluded.year,
            subject = excluded.subject,
            secretary = excluded.secretary,
            owner = excluded.owner,
            priority = excluded.priority,
            arrival_date = excluded.arrival_date,
            from_sector = excluded.from_sector,
            purpose = excluded.purpose,
            deadline = excluded.deadline,
            status = excluded.status,
            exit_date = excluded.exit_date,
            to_sector = excluded.to_sector,
            exit_purpose = excluded.exit_purpose,
            docs = excluded.docs,
            notes = excluded.notes,
            updated_at = excluded.updated_at
        """,
        values,
    )

    conn.execute("delete from history where process_id = ?", (process_id,))
    history_count = 0
    history_items = data_get(item, "history", default=[])
    if isinstance(history_items, list):
        for event in history_items:
            if not isinstance(event, dict):
                continue
            insert_history(
                conn,
                process_id=process_id,
                user_id=user_id,
                date=str(data_get(event, "date", default=data_get(item, "arrivalDate", "arrival_date", default=now()[:10]))),
                action=str(data_get(event, "action", default="Movimentacao registrada")),
                status_value=str(data_get(event, "status", default=status_value)),
                to=str(data_get(event, "to", "to_sector", default="")),
                purpose=str(data_get(event, "purpose", default="")),
                notes=str(data_get(event, "notes", default="")),
            )
            history_count += 1

    if history_count == 0:
        insert_history(
            conn,
            process_id=process_id,
            user_id=user_id,
            date=str(data_get(item, "arrivalDate", "arrival_date", default=now()[:10])),
            action="Entrada importada",
            status_value=status_value,
            to=str(data_get(item, "owner", default="")),
            purpose=str(data_get(item, "purpose", default="Importacao JSON")),
            notes="Processo importado para a base da API.",
        )
        history_count = 1

    return process_id, history_count


def init_db() -> None:
    legacy_db = LEGACY_DATA_DIR / "atlasflow.sqlite3"
    if not DB_PATH.exists() and legacy_db.exists():
        DB_PATH.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(legacy_db, DB_PATH)

    with db() as conn:
        conn.executescript(
            """
            create table if not exists users (
                id text primary key,
                name text not null,
                email text not null default '',
                department text not null default '',
                role text not null default 'Operador',
                password_hash text not null default '',
                password_salt text not null default '',
                created_at text not null
            );

            create table if not exists processes (
                id text primary key,
                user_id text not null references users(id) on delete cascade,
                number text not null default '',
                year text not null default '',
                subject text not null,
                secretary text not null default '',
                owner text not null default '',
                priority text not null default 'normal',
                arrival_date text not null default '',
                from_sector text not null default '',
                purpose text not null default '',
                deadline text not null default '',
                status text not null default 'entrada',
                exit_date text not null default '',
                to_sector text not null default '',
                exit_purpose text not null default '',
                docs text not null default '[]',
                notes text not null default '',
                created_at text not null,
                updated_at text not null
            );

            create table if not exists history (
                id text primary key,
                process_id text not null references processes(id) on delete cascade,
                user_id text not null references users(id) on delete cascade,
                date text not null,
                action text not null,
                status text not null,
                to_sector text not null default '',
                purpose text not null default '',
                notes text not null default '',
                created_at text not null
            );
            """
        )
        ensure_column(conn, "users", "password_hash", "text not null default ''")
        ensure_column(conn, "users", "password_salt", "text not null default ''")

        count = conn.execute("select count(*) as total from users").fetchone()["total"]
        if count == 0:
            created = now()
            conn.executemany(
                "insert into users (id, name, email, department, role, password_hash, password_salt, created_at) values (?, ?, ?, ?, ?, ?, ?, ?)",
                [
                    ("user-compras", "Israel Junior", "israel.junior@prefeitura.local", "Compras", "Administrador", "", "", created),
                    ("user-juridico", "Visitante", "visitante@prefeitura.local", "Visitante", "Consulta", "", "", created),
                ],
            )
        else:
            conn.execute(
                """
                update users
                   set name = 'Israel Junior',
                       email = case when email = 'compras@prefeitura.local' then 'israel.junior@prefeitura.local' else email end,
                       role = case when role = 'Operador' then 'Administrador' else role end
                 where id = 'user-compras'
                """
            )
            conn.execute(
                """
                update users
                   set name = 'Visitante',
                       email = case when email = 'juridico@prefeitura.local' then 'visitante@prefeitura.local' else email end,
                       department = case when department = 'Juridico' then 'Visitante' else department end,
                       role = case when role = 'Revisor' then 'Consulta' else role end
                 where id = 'user-juridico'
                """
            )


@app.on_event("startup")
def startup() -> None:
    init_db()


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "app": "AtlasFlow API"}


def clean_filename(value: str, extension: str) -> str:
    safe = "".join(char if char.isalnum() or char in ("-", "_") else "-" for char in value.lower())
    safe = "-".join(part for part in safe.split("-") if part)
    return f"{safe or 'documento'}.{extension}"


@app.post("/api/documents/export/docx")
def export_docx(payload: DocumentExport, authorization: str = Header(default="")) -> StreamingResponse:
    with db() as conn:
        require_session_user(conn, authorization)
    document = Document()
    document.add_heading(payload.title, level=1)
    for block in payload.text.split("\n"):
      if block.strip():
          document.add_paragraph(block)
      else:
          document.add_paragraph("")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    filename = clean_filename(payload.filename, "docx")
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/documents/export/pdf")
def export_pdf(payload: DocumentExport, authorization: str = Header(default="")) -> StreamingResponse:
    with db() as conn:
        require_session_user(conn, authorization)
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=48, leftMargin=48, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(payload.title, styles["Title"]), Spacer(1, 16)]
    for block in payload.text.split("\n"):
        story.append(Paragraph(block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or "&nbsp;", styles["BodyText"]))
        story.append(Spacer(1, 6))
    doc.build(story)
    output.seek(0)
    filename = clean_filename(payload.filename, "pdf")
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/users", response_model=list[User])
def list_users() -> list[dict[str, Any]]:
    with db() as conn:
        return [row_to_user(row) for row in conn.execute("select * from users order by name").fetchall()]


@app.post("/api/users", response_model=User, status_code=status.HTTP_201_CREATED)
def create_user(payload: UserCreate, admin_user_id: str = Query(...), authorization: str = Header(default="")) -> dict[str, Any]:
    user_id = new_id()
    salt, hashed = build_password_fields(payload.password)
    with db() as conn:
        require_admin_access(conn, admin_user_id, authorization)
        conn.execute(
            "insert into users (id, name, email, department, role, password_hash, password_salt, created_at) values (?, ?, ?, ?, ?, ?, ?, ?)",
            (user_id, payload.name.strip(), payload.email, payload.department, payload.role or "Operador", hashed, salt, now()),
        )
        return row_to_user(conn.execute("select * from users where id = ?", (user_id,)).fetchone())


@app.get("/api/users/{user_id}", response_model=User)
def get_user(user_id: str) -> dict[str, Any]:
    with db() as conn:
        return row_to_user(ensure_user(conn, user_id))


@app.patch("/api/users/{user_id}/password", response_model=User)
def update_user_password(user_id: str, payload: UserPasswordUpdate, authorization: str = Header(default="")) -> dict[str, Any]:
    salt, hashed = build_password_fields(payload.password)
    with db() as conn:
        require_admin_access(conn, payload.admin_user_id, authorization)
        ensure_user(conn, user_id)
        conn.execute(
            "update users set password_hash = ?, password_salt = ? where id = ?",
            (hashed, salt, user_id),
        )
        return row_to_user(conn.execute("select * from users where id = ?", (user_id,)).fetchone())


@app.delete("/api/users/{user_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_user(user_id: str, admin_user_id: str = Query(...), authorization: str = Header(default="")) -> None:
    if user_id == "user-compras":
        raise HTTPException(status_code=409, detail="O usuario administrador nao pode ser excluido")

    with db() as conn:
        require_admin_access(conn, admin_user_id, authorization)
        ensure_user(conn, user_id)
        process_count = conn.execute(
            "select count(*) as total from processes where user_id = ?",
            (user_id,),
        ).fetchone()["total"]
        if process_count:
            raise HTTPException(status_code=409, detail="Usuario possui processos vinculados")
        conn.execute("delete from users where id = ?", (user_id,))


@app.post("/api/import", response_model=ImportResult)
def import_backup(payload: ImportPayload, authorization: str = Header(default="")) -> dict[str, Any]:
    current_user_id = payload.currentUserId or payload.current_user_id
    with db() as conn:
        session_user = require_session_user(conn, authorization)
        if not is_admin_row(session_user):
            raise HTTPException(status_code=403, detail="Apenas administradores podem importar dados")
        user_ids = [upsert_import_user(conn, item) for item in payload.users if isinstance(item, dict)]
        fallback_user_id = current_user_id or (user_ids[0] if user_ids else "user-compras")
        if not conn.execute("select 1 from users where id = ?", (fallback_user_id,)).fetchone():
            upsert_import_user(conn, {"id": fallback_user_id, "name": "Usuario importado"})

        history_total = 0
        process_total = 0
        for item in payload.processes:
            if not isinstance(item, dict):
                continue
            _, history_count = upsert_import_process(conn, item, fallback_user_id)
            history_total += history_count
            process_total += 1

        return {
            "users": len(user_ids),
            "processes": process_total,
            "history": history_total,
            "current_user_id": fallback_user_id,
        }


@app.post("/api/auth/login", response_model=AuthSession)
def login(payload: AuthLogin) -> dict[str, Any]:
    with db() as conn:
        user = ensure_user(conn, payload.user_id)
        saved_hash = user["password_hash"]
        saved_salt = user["password_salt"]
        if REQUIRE_PASSWORDS and not saved_hash:
            raise HTTPException(status_code=401, detail="Usuario precisa ter senha configurada antes de entrar")
        if saved_hash and password_hash(payload.password, saved_salt) != saved_hash:
            raise HTTPException(status_code=401, detail="Senha invalida")
        if not saved_hash and payload.password:
            raise HTTPException(status_code=401, detail="Usuario ainda nao possui senha configurada")

        token_seed = f"{user['id']}:{now()}:{secrets.token_hex(16)}"
        token = hashlib.sha256(token_seed.encode("utf-8")).hexdigest()
        SESSIONS[token] = user["id"]
        return {
            "user_id": user["id"],
            "name": user["name"],
            "department": user["department"],
            "role": user["role"],
            "token": token,
        }


@app.get("/api/users/{user_id}/processes", response_model=list[Process])
def list_user_processes(user_id: str, authorization: str = Header(default="")) -> list[dict[str, Any]]:
    with db() as conn:
        require_user_access(conn, user_id, authorization)
        rows = conn.execute(
            "select * from processes where user_id = ? order by coalesce(nullif(deadline, ''), '9999-12-31'), created_at",
            (user_id,),
        ).fetchall()
        return [row_to_process(row) for row in rows]


@app.post("/api/users/{user_id}/processes", response_model=ProcessWithHistory, status_code=status.HTTP_201_CREATED)
def create_process(user_id: str, payload: ProcessCreate, authorization: str = Header(default="")) -> dict[str, Any]:
    validate_status(payload.status)
    validate_priority(payload.priority)
    process_id = new_id()
    created = now()
    with db() as conn:
        session_user = require_user_access(conn, user_id, authorization)
        if is_read_only_row(session_user):
            raise HTTPException(status_code=403, detail="Perfil permite apenas consulta")
        conn.execute(
            """
            insert into processes (
                id, user_id, number, year, subject, secretary, owner, priority, arrival_date,
                from_sector, purpose, deadline, status, exit_date, to_sector, exit_purpose,
                docs, notes, created_at, updated_at
            )
            values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                process_id,
                user_id,
                payload.number,
                payload.year,
                payload.subject,
                payload.secretary,
                payload.owner,
                payload.priority,
                payload.arrival_date,
                payload.from_sector,
                payload.purpose,
                payload.deadline,
                payload.status,
                payload.exit_date,
                payload.to_sector,
                payload.exit_purpose,
                json.dumps(payload.docs, ensure_ascii=False),
                payload.notes,
                created,
                created,
            ),
        )
        insert_history(
            conn,
            process_id=process_id,
            user_id=user_id,
            date=payload.arrival_date or created[:10],
            action="Entrada registrada",
            status_value=payload.status,
            to=payload.owner,
            purpose=payload.purpose or "Cadastro inicial",
            notes="Processo cadastrado na API AtlasFlow.",
        )
        return get_process_payload(conn, process_id)


@app.get("/api/processes/{process_id}", response_model=ProcessWithHistory)
def get_process(process_id: str, authorization: str = Header(default="")) -> dict[str, Any]:
    with db() as conn:
        process = ensure_process(conn, process_id)
        require_process_access(conn, process, authorization)
        return get_process_payload(conn, process_id)


@app.put("/api/processes/{process_id}", response_model=ProcessWithHistory)
def update_process(process_id: str, payload: ProcessUpdate, authorization: str = Header(default="")) -> dict[str, Any]:
    validate_status(payload.status)
    validate_priority(payload.priority)
    with db() as conn:
        old = ensure_process(conn, process_id)
        require_process_write(conn, old, authorization)
        conn.execute(
            """
            update processes
               set number = ?, year = ?, subject = ?, secretary = ?, owner = ?, priority = ?,
                   arrival_date = ?, from_sector = ?, purpose = ?, deadline = ?, status = ?,
                   exit_date = ?, to_sector = ?, exit_purpose = ?, docs = ?, notes = ?, updated_at = ?
             where id = ?
            """,
            (
                payload.number,
                payload.year,
                payload.subject,
                payload.secretary,
                payload.owner,
                payload.priority,
                payload.arrival_date,
                payload.from_sector,
                payload.purpose,
                payload.deadline,
                payload.status,
                payload.exit_date,
                payload.to_sector,
                payload.exit_purpose,
                json.dumps(payload.docs, ensure_ascii=False),
                payload.notes,
                now(),
                process_id,
            ),
        )
        if old["status"] != payload.status:
            insert_history(
                conn,
                process_id=process_id,
                user_id=old["user_id"],
                date=now()[:10],
                action="Status alterado",
                status_value=payload.status,
                to=payload.to_sector,
                purpose=payload.exit_purpose or payload.purpose,
                notes="Alteracao feita pela edicao do cadastro.",
            )
        return get_process_payload(conn, process_id)


@app.delete("/api/processes/{process_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_process(process_id: str, authorization: str = Header(default="")) -> None:
    with db() as conn:
        process = ensure_process(conn, process_id)
        require_process_write(conn, process, authorization)
        conn.execute("delete from processes where id = ?", (process_id,))


@app.post("/api/processes/{process_id}/movements", response_model=ProcessWithHistory)
def move_process(process_id: str, payload: MovementCreate, authorization: str = Header(default="")) -> dict[str, Any]:
    validate_status(payload.status)
    with db() as conn:
        process = ensure_process(conn, process_id)
        require_process_write(conn, process, authorization)
        exit_date = payload.date if payload.status in {"devolver", "concluido"} else process["exit_date"]
        to_sector = payload.to if payload.status in {"devolver", "concluido"} else process["to_sector"]
        exit_purpose = payload.purpose if payload.status in {"devolver", "concluido"} else process["exit_purpose"]
        conn.execute(
            """
            update processes
               set status = ?, exit_date = ?, to_sector = ?, exit_purpose = ?, updated_at = ?
             where id = ?
            """,
            (payload.status, exit_date, to_sector, exit_purpose, now(), process_id),
        )
        insert_history(
            conn,
            process_id=process_id,
            user_id=process["user_id"],
            date=payload.date or now()[:10],
            action="Movimentacao registrada",
            status_value=payload.status,
            to=payload.to,
            purpose=payload.purpose,
            notes=payload.notes,
        )
        return get_process_payload(conn, process_id)


@app.get("/api/users/{user_id}/history", response_model=list[HistoryItem])
def list_user_history(user_id: str, authorization: str = Header(default="")) -> list[dict[str, Any]]:
    with db() as conn:
        require_user_access(conn, user_id, authorization)
        rows = conn.execute(
            "select * from history where user_id = ? order by date desc, created_at desc",
            (user_id,),
        ).fetchall()
        return [row_to_history(row) for row in rows]


def get_process_payload(conn: sqlite3.Connection, process_id: str) -> dict[str, Any]:
    process = row_to_process(ensure_process(conn, process_id))
    history = conn.execute(
        "select * from history where process_id = ? order by date desc, created_at desc",
        (process_id,),
    ).fetchall()
    process["history"] = [row_to_history(row) for row in history]
    return process
